import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from datetime import datetime
from num2words import num2words
from db.connection import get_connection  # Use shared connection file
import os

# Get next invoice number
def get_next_invoice_number():
    conn = get_connection()
    c = conn.cursor()
    c.execute("SELECT MAX(invoice_no) FROM invoices")
    result = c.fetchone()
    conn.close()
    return (result[0] or 0) + 1

def reset_form():
    invoice_no_var.set(get_next_invoice_number())
    date_var.set(datetime.today().strftime("%d-%m-%Y"))
    buyer_var.set("")
    buyer_address_var.set("")
    consignee_var.set("")
    consignee_address_var.set("")
    vehicle_no_var.set("")
    destination_var.set("")
    quantity_var.set("")
    rate_var.set("")
    bags_var.set("")
    total_amount_label.config(text="0.00")
    amount_words_var.set("")
    bank_var.set(list(bank_details.keys())[0])
    update_bank_info()

def close_app():
    root.destroy()

# Save to DB
def save_invoice_to_db():
    conn = get_connection()
    c = conn.cursor()
    c.execute("""
        INSERT INTO invoices (
            invoice_no, date, buyer_name, buyer_address, consignee, consignee_address,
            vehicle_no, destination, quantity, rate, total_amount, bank, account_no, ifsc,
            amount_words, bags
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """, (
        invoice_no_var.get(),
        date_var.get(),
        buyer_var.get(),
        buyer_address_var.get(),
        consignee_var.get(),
        consignee_address_var.get(),
        vehicle_no_var.get(),
        destination_var.get(),
        quantity_var.get(),
        rate_var.get(),
        float(total_amount_label["text"]),
        bank_var.get(),
        account_no_var.get(),
        ifsc_var.get(),
        amount_words_var.get(),
        bags_var.get()
    ))
    conn.commit()
    conn.close()

# Bank Details
bank_details = {
    "SBI": {"account_no": "1234567890", "ifsc": "SBIN0001234"},
    "PNB": {"account_no": "9876543210", "ifsc": "PUNB0123456"}
}

def update_bank_info(event=None):
    selected = bank_var.get()
    details = bank_details.get(selected, {})
    account_no_var.set(details.get("account_no", ""))
    ifsc_var.set(details.get("ifsc", ""))

def update_total(*args):
    try:
        qty = float(quantity_var.get())
        rate = float(rate_var.get())
        total = (qty / 100) * rate
        total_amount_label.config(text=f"{total:.2f}")

        rupees = int(total)
        paise = int(round((total - rupees) * 100))
        if paise > 0:
            words = f"{num2words(rupees, lang='en_IN').title()} Rupees and {num2words(paise, lang='en_IN').title()} Paise Only"
        else:
            words = f"{num2words(rupees, lang='en_IN').title()} Rupees Only"

        amount_words_var.set(words)
    except ValueError:
        total_amount_label.config(text="0.00")
        amount_words_var.set("")

def replace_placeholder(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                inline_text = paragraph.text.replace(key, value.strip())
                for run in paragraph.runs:
                    run.clear()
                paragraph.text = inline_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                for key, value in replacements.items():
                    if key in cell_text:
                        cell.text = cell_text.replace(key, value.strip())

def generate_invoice():
    summary = f"""
    Invoice No: {invoice_no_var.get()}
    Date: {date_var.get()}

    Buyer: {buyer_var.get().upper()}
    Buyer Address: {buyer_address_var.get()}

    Consignee: {consignee_var.get().upper()}
    Consignee Address: {consignee_address_var.get()}

    Vehicle No: {vehicle_no_var.get()}
    Destination: {destination_var.get()}
    Bags: {bags_var.get()}

    Quantity: {quantity_var.get()} kg
    Rate: ₹{rate_var.get()} per quintal
    Total: ₹{total_amount_label['text']}
    Amount in Words: {amount_words_var.get()}

    Bank: {bank_var.get()}
    Account No: {account_no_var.get()}
    IFSC: {ifsc_var.get()}
    """

    confirm = messagebox.askyesno("Confirm Invoice Details", summary + "Do you want to generate the invoice?")
    if not confirm:
        return

    save_invoice_to_db()
    doc = Document("templates/invoice.docx")
    inline_replacements = {
        '{{invoice_no}}': str(invoice_no_var.get()),
        '{{date}}': date_var.get(),
        '{{buyer_full}}': f"{buyer_var.get().upper()}\n{buyer_address_var.get()}",
        '{{consignee_full}}': f"{consignee_var.get().upper()}\n{consignee_address_var.get()}",
        '{{vehicle_no}}': vehicle_no_var.get(),
        '{{destination}}': destination_var.get(),
        '{{quantity}}': f"{quantity_var.get()} kg",
        '{{rate}}': f"{rate_var.get()} per quintal",
        '{{total}}': total_amount_label['text'],
        '{{bank}}': bank_var.get(),
        '{{account_no}}': account_no_var.get(),
        '{{ifsc}}': ifsc_var.get(),
        '{{amount_words}}': amount_words_var.get(),
        '{{bags}}': bags_var.get()
    }
    replace_placeholder(doc, inline_replacements)

    save_dir = os.path.join(os.getcwd(), "bills 2025")
    os.makedirs(save_dir, exist_ok=True)

    vehicle_clean = vehicle_no_var.get().strip().upper().replace(" ", "").replace("-", "")
    vehicle_parts = [vehicle_clean[:2], vehicle_clean[2:4], vehicle_clean[4:6], vehicle_clean[6:]]
    filename = os.path.join(
        save_dir,
        f"{buyer_var.get().strip().replace(' ', '_')}_{date_var.get().replace('/', '-')}" +
        f"_{'-'.join(vehicle_parts)}.docx"
    )
    doc.save(filename)
    messagebox.showinfo("Success", f"Invoice {filename} generated successfully!")
# GUI
root = tk.Tk()
root.title(" Maize Merchant- Invoice Generator")
root.geometry("900x750")
root.configure(bg="#f2f7fb")

heading = tk.Label(root, text=" Maize Merchant", font=("Helvetica", 18, "bold"), bg="#f2f7fb", fg="#234e70")
heading.pack(pady=10)

frame = tk.Frame(root, bg="#f2f7fb")
frame.pack(pady=10)

# Variables
invoice_no_var = tk.IntVar(value=get_next_invoice_number())
date_var = tk.StringVar(value=datetime.today().strftime("%d-%m-%Y"))
buyer_var = tk.StringVar()
buyer_address_var = tk.StringVar()
consignee_var = tk.StringVar()
consignee_address_var = tk.StringVar()
vehicle_no_var = tk.StringVar()
destination_var = tk.StringVar()
quantity_var = tk.StringVar()
rate_var = tk.StringVar()
bank_var = tk.StringVar()
account_no_var = tk.StringVar()
ifsc_var = tk.StringVar()
amount_words_var = tk.StringVar()
bags_var = tk.StringVar()

# Section: Invoice Info
sec1 = tk.LabelFrame(frame, text="Invoice Details", font=("Arial", 10, "bold"), bg="#f2f7fb")
sec1.grid(row=0, column=0, columnspan=2, pady=5, sticky="ew")
tk.Label(sec1, text="Invoice No:", bg="#f2f7fb").grid(row=0, column=0, padx=5, pady=5)
tk.Entry(sec1, textvariable=invoice_no_var, width=10).grid(row=0, column=1)
tk.Label(sec1, text="Date:", bg="#f2f7fb").grid(row=0, column=2)
tk.Entry(sec1, textvariable=date_var, width=15).grid(row=0, column=3)

# Section: Buyer
sec2 = tk.LabelFrame(frame, text="Buyer & Consignee Info", font=("Arial", 10, "bold"), bg="#f2f7fb")
sec2.grid(row=1, column=0, pady=5, sticky="ew")
tk.Label(sec2, text="Buyer Name:", bg="#f2f7fb").grid(row=0, column=0, padx=5, pady=5)
tk.Entry(sec2, textvariable=buyer_var).grid(row=0, column=1)
tk.Label(sec2, text="Buyer Address:", bg="#f2f7fb").grid(row=1, column=0, padx=5, pady=5)
tk.Entry(sec2, textvariable=buyer_address_var).grid(row=1, column=1)

tk.Label(sec2, text="Consignee Name:", bg="#f2f7fb").grid(row=2, column=0, padx=5, pady=5)
tk.Entry(sec2, textvariable=consignee_var).grid(row=2, column=1)
tk.Label(sec2, text="Consignee Address:", bg="#f2f7fb").grid(row=3, column=0, padx=5, pady=5)
tk.Entry(sec2, textvariable=consignee_address_var).grid(row=3, column=1)

# Section: Shipping
sec3 = tk.LabelFrame(frame, text="Shipment Details", font=("Arial", 10, "bold"), bg="#f2f7fb")
sec3.grid(row=1, column=1, pady=5, sticky="ew")
tk.Label(sec3, text="Vehicle No:", bg="#f2f7fb").grid(row=0, column=0, padx=5, pady=5)
tk.Entry(sec3, textvariable=vehicle_no_var).grid(row=0, column=1)
tk.Label(sec3, text="Destination:", bg="#f2f7fb").grid(row=1, column=0, padx=5, pady=5)
tk.Entry(sec3, textvariable=destination_var).grid(row=1, column=1)
tk.Label(sec3, text="Bags:", bg="#f2f7fb").grid(row=2, column=0, padx=5, pady=5)
tk.Entry(sec3, textvariable=bags_var).grid(row=2, column=1)

# Section: Transaction
sec4 = tk.LabelFrame(frame, text="Transaction", font=("Arial", 10, "bold"), bg="#f2f7fb")
sec4.grid(row=2, column=0, columnspan=2, pady=5, sticky="ew")
tk.Label(sec4, text="Quantity (kg):", bg="#f2f7fb").grid(row=0, column=0)
tk.Entry(sec4, textvariable=quantity_var).grid(row=0, column=1)
tk.Label(sec4, text="Rate (per quintal):", bg="#f2f7fb").grid(row=0, column=2)
tk.Entry(sec4, textvariable=rate_var).grid(row=0, column=3)
tk.Label(sec4, text="Total Amount:", bg="#f2f7fb").grid(row=0, column=4)
total_amount_label = tk.Label(sec4, text="0.00", font=("Arial", 10, "bold"), bg="#f2f7fb")
total_amount_label.grid(row=0, column=5, padx=10)

# Section: Bank
sec5 = tk.LabelFrame(frame, text="Bank Details", font=("Arial", 10, "bold"), bg="#f2f7fb")
sec5.grid(row=3, column=0, columnspan=2, pady=5, sticky="ew")
tk.Label(sec5, text="Select Bank:", bg="#f2f7fb").grid(row=0, column=0)
bank_dropdown = ttk.Combobox(sec5, textvariable=bank_var, values=list(bank_details.keys()))
bank_dropdown.grid(row=0, column=1)
bank_dropdown.bind("<<ComboboxSelected>>", update_bank_info)
tk.Label(sec5, text="Account No:", bg="#f2f7fb").grid(row=0, column=2)
tk.Entry(sec5, textvariable=account_no_var).grid(row=0, column=3)
tk.Label(sec5, text="IFSC Code:", bg="#f2f7fb").grid(row=0, column=4)
tk.Entry(sec5, textvariable=ifsc_var).grid(row=0, column=5)

# Amount in words
tk.Label(frame, text="Amount in Words:", font=("Arial", 10, "bold"), bg="#f2f7fb").grid(row=4, column=0, pady=10)
tk.Entry(frame, textvariable=amount_words_var, width=80).grid(row=4, column=1)

# Button
tk.Button(root, text="Generate Invoice", font=("Arial", 10, "bold"), command=generate_invoice).pack(pady=20)

# Add a Reset/New Bill button
btn_reset = tk.Button(root, text="New Bill", command=reset_form, bg="orange", fg="white")
btn_reset.pack(pady=5)

# Add a Close button
btn_close = tk.Button(root, text="Close App", command=close_app, bg="#dc3545", fg="white", font=("Helvetica", 10))
btn_close.pack(pady=5)

# Bind total calculation
quantity_var.trace("w", update_total)
rate_var.trace("w", update_total)

root.mainloop()
